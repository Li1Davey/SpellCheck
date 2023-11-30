'''
Program by Cristian Pedroza & David Sanchez

code for reading and creating a copy of a document from
https://stackoverflow.com/questions/48869423/how-do-i-copy-the-contents-of-a-word-document

'''
import docx

# Function to read a Word document and extract text
def read_word_document(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

# Function to write the corrected text to a new Word document
def write_corrected_text_to_docx(corrected_text, output_path='corrected_document.docx'):
    doc = docx.Document()
    paragraphs = corrected_text.split('\n')
    
    for paragraph_text in paragraphs:
        doc.add_paragraph(paragraph_text)
    
    doc.save(output_path)
